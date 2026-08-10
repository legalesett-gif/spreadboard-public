#!/usr/bin/env python3
"""Build the send-ready SpreadBoard affiliate partner pack.

Design system:
- preset: launch_messaging_guide (compact_reference_guide)
- first-page pattern: customer_pack
- named brand override: SpreadBoard teal headings/accent and pale mint fills
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "SpreadBoard Affiliate Partner Pack.docx"

INK = "102A2A"
TEAL = "0F766E"
TEAL_DARK = "115E59"
MINT = "E8F5F2"
MINT_STRONG = "D5EFE9"
MUTED = "526463"
PALE = "F4F8F7"
WHITE = "FFFFFF"
LINE = "B8D4CF"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run(run, *, size=11, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color=LINE, size=6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int]) -> None:
    assert sum(widths) == CONTENT_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, *, color=LINE, size=6, space=6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_numbering(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def body_para(doc, text: str, *, bold_lead: str | None = None, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True)
        set_run(p.add_run(text[len(bold_lead):]), italic=italic)
    else:
        set_run(p.add_run(text), italic=italic)
    return p


def bullet(doc, text: str, num_id: int, *, bold_lead: str | None = None):
    p = body_para(doc, text, bold_lead=bold_lead, after=4)
    apply_num(p, num_id)
    return p


def numbered(doc, text: str, num_id: int, *, bold_lead: str | None = None):
    p = body_para(doc, text, bold_lead=bold_lead, after=4)
    apply_num(p, num_id)
    return p


def h1(doc, text: str):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size=16, bold=True, color=TEAL_DARK)
    return p


def h2(doc, text: str):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size=13, bold=True, color=TEAL)
    return p


def table_text(cell, text: str, *, bold=False, color=INK, size=9.5, align=None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), bold=bold, color=color, size=size)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], *, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        set_cell_fill(table.rows[0].cells[index], MINT_STRONG)
        table_text(table.rows[0].cells[index], header, bold=True, color=TEAL_DARK, size=9)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_fill(cells[index], WHITE)
            table_text(
                cells[index], value, size=9.2,
                align=WD_ALIGN_PARAGRAPH.RIGHT if index in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT,
            )
        set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL_DARK, 18, 10),
        ("Heading 2", 13, TEAL, 14, 7),
        ("Heading 3", 12, TEAL_DARK, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("SpreadBoard  |  Affiliate Partner Pack"), size=9, bold=True, color=TEAL_DARK)
    set_paragraph_border(p)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run("PARTNER ENABLEMENT"), size=9, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run("SpreadBoard Affiliate Partner Pack"), size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run("A recurring-revenue partnership for crypto research audiences"), size=13.5, color=MUTED)

    body_para(
        doc,
        "SpreadBoard helps crypto-market participants compare live cross-venue spreads, funding carry, fair value, transfer conditions, and route history in one research workspace. Approved creators receive a personal referral link, a private tracking cabinet, and recurring commission on referred memberships.",
        after=14,
    )

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    metrics = [
        ("VIEWER OFFER", "20% off", "first 30-day value"),
        ("PARTNER", "50%", "recurring commission"),
        ("ATTRIBUTION", "90 days", "then fixed at registration"),
        ("PAYOUT", "Weekly", "USDT on Arbitrum"),
    ]
    for cell, (label, value, detail) in zip(table.rows[0].cells, metrics):
        set_cell_fill(cell, MINT)
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_after = Pt(3)
        set_run(p1.add_run(label), size=8, bold=True, color=TEAL)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        set_run(p2.add_run(value), size=16, bold=True, color=INK)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        set_run(p3.add_run(detail), size=8.5, color=MUTED)
    set_table_geometry(table, [2340, 2340, 2340, 2340])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Program overview  •  10 August 2026  •  spreadarbitrage.ink"), size=9, color=MUTED)


def build() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    bullets = add_numbering(doc, ordered=False)
    numbers = add_numbering(doc, ordered=True)
    add_cover(doc)

    h1(doc, "What the audience receives")
    body_para(
        doc,
        "SpreadBoard is a read-only research service built around public market data. It does not hold customer funds, connect to customer exchange accounts, or execute trades.",
    )
    for lead, text in (
        ("Live arbitrage discovery.", " Live cross-venue routes with market type, direction, matched depth, freshness, and transfer context."),
        ("Funding research.", " Normalized carry comparisons, settled and projected funding context, and multi-period history."),
        ("Evidence workspace.", " Custom spread charts, fair-price views, market events, watchlists, alerts, and position journaling."),
        ("Two membership choices.", " Scanner for live discovery; Research Pro for the full research workspace and private Telegram forum."),
    ):
        bullet(doc, lead + text, bullets, bold_lead=lead)
    body_para(
        doc,
        "The product is intended for users who understand that spreads can close before entry, funding can reverse, liquidity can disappear, transfers can be restricted, and market data may be delayed or unavailable. It is research data, not investment advice and not a promise of profit.",
        italic=True,
    )

    h1(doc, "How the partnership works")
    steps = [
        ("Cabinet setup.", " SpreadBoard creates the partner account and generates a readable referral link from the channel name. The link label does not affect attribution or commission."),
        ("Audience click.", " A qualifying click is remembered for up to 90 days."),
        ("Registration.", " The first qualifying referral is fixed when the user registers, preserving future renewal attribution."),
        ("First purchase.", " The customer receives 20% off one 30-day membership value. A longer first term receives the same one-month discount once."),
        ("Commission.", " Each confirmed crypto invoice creates commission equal to 50% of settled plan revenue after the discount. Invoice-identification cents, refunds, reversals, fraud, and duplicate or self-referred accounts are excluded."),
        ("Weekly payout.", " Commission is held for seven days, then included in the next manual payout batch and sent in USDT on Arbitrum. The transaction reference remains in the cabinet."),
    ]
    for lead, text in steps:
        numbered(doc, lead + text, numbers, bold_lead=lead)

    h1(doc, "Illustrative 30-day economics")
    add_table(
        doc,
        ["Membership", "List price", "Customer first 30 days", "First commission", "Later 30-day commission"],
        [
            ["Scanner", "$49.00", "$39.20", "$19.60", "$24.50"],
            ["Research Pro", "$149.00", "$119.20", "$59.60", "$74.50"],
        ],
        [1740, 1300, 2240, 1880, 2200],
        numeric_cols=(1, 2, 3, 4),
    )
    body_para(
        doc,
        "Examples assume a confirmed 30-day invoice and no refund, reversal, tax, or adjustment. For a 90- or 365-day first invoice, only one 30-day membership value receives the 20% discount; commission is then 50% of the remaining settled plan revenue.",
        italic=True,
    )

    doc.add_page_break()
    h1(doc, "What the partner cabinet shows")
    for text in (
        "Personal referral link and copy button",
        "Qualifying link visits and attributed registrations",
        "Paying-customer count",
        "Membership tier and term for each settled invoice",
        "Net settled subscription revenue and the partner's 50% commission",
        "Commission status: on hold, payable, in a batch, paid, or void with a reason",
        "Payable balance, paid-to-date total, and payout-batch history",
        "Saved USDT-on-Arbitrum payout destination and payment reference",
    ):
        bullet(doc, text, bullets)
    body_para(
        doc,
        "Partners do not receive customer passwords, wallet credentials, exchange credentials, or customer email addresses.",
        italic=True,
    )

    h1(doc, "Clear, compliant promotion")
    body_para(
        doc,
        "Creators should explain what they actually reviewed and distinguish the research product from trading execution. Every promotion must disclose the paid affiliate relationship clearly and conspicuously.",
    )
    h2(doc, "Required disclosure for YouTube")
    body_para(doc, "Say and show near the start of the promotion, and repeat immediately before the link in the description:")
    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    set_cell_fill(callout.cell(0, 0), MINT)
    table_text(
        callout.cell(0, 0),
        "AD — paid affiliate promotion. I receive a commission if you subscribe through this link.",
        bold=True, color=TEAL_DARK, size=11,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("Using YouTube's paid-promotion tool is encouraged, but does not replace the spoken, on-screen, and description disclosure above."), size=9.5, italic=True, color=MUTED)

    h2(doc, "Good content angles")
    for text in (
        "Demonstrate how you verify a spread rather than quoting a headline number.",
        "Compare funding carry across venues and discuss reversal risk.",
        "Show the evidence behind a route: market type, direction, depth, freshness, transfer status, and history.",
        "Explain which membership tier suits a research-only user.",
        "State that viewers must do their own research and manage their own execution and risk.",
    ):
        bullet(doc, text, bullets)

    h2(doc, "Claims that are not permitted")
    for text in (
        "Guaranteed, risk-free, or certain profits",
        "Specific future returns or a claim that historical figures will repeat",
        "A claim that SpreadBoard executes, advises on, or safeguards trades or funds",
        "Fabricated results, undisclosed endorsements, spam, cookie stuffing, self-referrals, fake engagement, or bidding on the SpreadBoard brand",
    ):
        bullet(doc, text, bullets)

    h1(doc, "Onboarding checklist")
    body_para(doc, "To create the cabinet and agreement, SpreadBoard needs:")
    for text in (
        "Channel or public name and channel URL",
        "Partner's legal name or business entity",
        "Country and business or contact address",
        "Login and notice email",
        "Primary audience countries",
        "USDT wallet address on Arbitrum; this can be added securely in the cabinet",
        "Confirmation that the affiliate disclosure and program terms are accepted",
    ):
        bullet(doc, text, bullets)
    body_para(
        doc,
        "There is no minimum sales requirement and no exclusivity unless separately agreed in writing. The final signed Affiliate Partner Agreement and the current Affiliate Program Terms govern the relationship.",
    )

    h1(doc, "Next step")
    body_para(
        doc,
        "Reply to the invitation email confirming interest. SpreadBoard will send a single-use account setup link so the partner chooses their own password, followed by the personal referral link and agreement for signature.",
    )
    info = doc.add_table(rows=2, cols=2)
    set_table_geometry(info, [2700, 6660])
    for row, (label, value) in zip(info.rows, (
        ("Website", "https://spreadarbitrage.ink"),
        ("Affiliate terms", "https://spreadarbitrage.ink/affiliate-terms"),
    )):
        set_cell_fill(row.cells[0], MINT)
        set_cell_fill(row.cells[1], WHITE)
        table_text(row.cells[0], label, bold=True, color=TEAL_DARK)
        table_text(row.cells[1], value)
    set_table_geometry(info, [2700, 6660])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(
        p.add_run("Program overview dated 10 August 2026. Commercial terms may be updated only as allowed by the signed agreement. SpreadBoard is a research product, not investment advice."),
        size=8.5, italic=True, color=MUTED,
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
